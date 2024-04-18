import os
from docx import Document
import random
import jieba
import argparse
from nltk.tokenize import word_tokenize, sent_tokenize


def read_and_analyze_file(filepath, lang):
    total_lines = 0
    total_words = 0
    unique_words = set()

    with open(filepath, "r", encoding="utf-8") as file:
        for line in file:
            total_lines += 1
            words = word_tokenize(line) if lang != "zh" else list(jieba.cut(line))
            total_words += len(words)
            unique_words.update(words)

    return total_lines, total_words, len(unique_words)


def extract_and_divide_into_docx(
    filepath, output_path, output_basename, max_words_per_doc, lang
):
    document = Document()
    word_count = 0
    doc_count = 0
    line_count = 0  # Initialize line counter

    # Create the temporary path within the output folder if it doesn't exist
    temp_folder_path = os.path.join(os.path.dirname(filepath), output_path)

    if not os.path.exists(temp_folder_path):
        os.makedirs(temp_folder_path)

    with open(filepath, "r", encoding="utf-8") as file:
        for line in file:
            line_count += 1  # Increment line counter for each line processed
            words = word_tokenize(line) if lang != "zh" else list(jieba.cut(line))
            word_count += len(words)
            if (
                word_count > max_words_per_doc or line_count == 1
            ):  # Check if new document is needed or it's the first line
                if doc_count > 0 or line_count > 1:
                    output_filename = f"{output_basename}_{doc_count}.docx"
                    document.save(os.path.join(temp_folder_path, output_filename))

                    print(
                        f"Saved {output_filename} with {word_count - len(words)} words and {line_count - 1} lines."
                    )

                # Reset counters and document for new file
                document = Document()
                word_count = len(words)
                doc_count += 1
                line_count = 1  # Reset line counter for the new document

            document.add_paragraph(line)

        # Save the last document
        if word_count > 0:
            output_filename = f"{output_basename}_{doc_count}.docx"

            document.save(os.path.join(temp_folder_path, output_filename))
            print(
                f"Saved {output_filename} with {word_count} words and {line_count} lines."
            )


def main():

    parser = argparse.ArgumentParser(
        description="splitting word documents"
    )

    parser.add_argument(
        "--language", type=str, required=True, help="the language of the file"
    )

    parser.add_argument(
        "--folder_path", type=str, required=True, help="the folder where file is located in"
    )

    parser.add_argument(
        "--base_filename", type=str, required=True, help="the name of the file"
    )

    parser.add_argument(
        "--output_path", type=str, required=True, help="the folder for all the output files"
    )

    parser.add_argument(
        "--output_basename", type=str, required=True, help="the name for all the output files"
    )

    args = parser.parse_args()

    filepath = f"{args.folder_path}/{args.base_filename}.txt"

    max_words_per_doc = random.randint(56789, 60606)
    # max_words_per_doc = random.randint(12345, 60606)

    # Analyze the file (optional, can be removed if not needed)
    total_lines, total_words, unique_word_count = read_and_analyze_file(
        filepath, args.language
    )
    print(
        f"Total Lines: {total_lines}, Total Words: {total_words}, Unique Words: {unique_word_count}"
    )

    # Extract lines and divide them into multiple docx files
    extract_and_divide_into_docx(
        filepath, args.output_path, args.output_basename, max_words_per_doc, args.language
    )

    

if __name__ == "__main__":
    main()
